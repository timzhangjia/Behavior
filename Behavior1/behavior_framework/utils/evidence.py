"""
Test Evidence Manager - Saves test execution evidence (API requests, responses, DB queries, etc.)
"""

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List
from behavior_framework.utils.logger import Logger

logger = Logger()

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Word document generation will be disabled.")


class EvidenceManager:
    """Manager for saving test execution evidence"""
    
    def __init__(self, evidence_dir: str = "evidence"):
        """
        Initialize Evidence Manager
        
        Args:
            evidence_dir: Directory to save evidence files
        """
        self.evidence_dir = Path(evidence_dir)
        self.evidence_dir.mkdir(parents=True, exist_ok=True)
        self.current_evidence: Dict[str, Any] = {}
    
    def start_scenario(self, feature_name: str, scenario_name: str):
        """
        Initialize evidence for a new scenario
        
        Args:
            feature_name: Feature name
            scenario_name: Scenario name
        """
        # Sanitize names for file system
        safe_feature = self._sanitize_filename(feature_name)
        safe_scenario = self._sanitize_filename(scenario_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        self.current_evidence = {
            "feature": feature_name,
            "scenario": scenario_name,
            "timestamp": timestamp,
            "api_requests": [],
            "database_queries": [],
            "ui_screenshots": []  # List of screenshot paths
        }
        
        # Create temporary directory for screenshots
        self.screenshots_dir = self.evidence_dir / "temp_screenshots" / f"{safe_feature}_{safe_scenario}_{timestamp}"
        self.screenshots_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Evidence collection started for scenario: {scenario_name}")
    
    def add_api_request(
        self,
        method: str,
        url: str,
        headers: Optional[Dict[str, str]] = None,
        body: Optional[Any] = None,
        response_status: Optional[int] = None,
        response_headers: Optional[Dict[str, str]] = None,
        response_body: Optional[Any] = None
    ):
        """
        Add API request evidence
        
        Args:
            method: HTTP method (GET, POST, etc.)
            url: Request URL
            headers: Request headers
            body: Request body
            response_status: Response status code
            response_headers: Response headers
            response_body: Response body
        """
        request_evidence = {
            "timestamp": datetime.now().isoformat(),
            "request": {
                "method": method,
                "url": url,
                "headers": headers or {},
                "body": self._serialize_body(body)
            },
            "response": {
                "status_code": response_status,
                "headers": response_headers or {},
                "body": self._serialize_body(response_body)
            }
        }
        
        self.current_evidence["api_requests"].append(request_evidence)
        logger.debug(f"API request evidence added: {method} {url}")
    
    def add_database_query(
        self,
        query: str,
        result: Optional[List[Dict[str, Any]]] = None,
        error: Optional[str] = None
    ):
        """
        Add database query evidence
        
        Args:
            query: SQL query
            result: Query result
            error: Error message if query failed
        """
        query_evidence = {
            "timestamp": datetime.now().isoformat(),
            "query": query,
            "result": result or [],
            "row_count": len(result) if result else 0,
            "error": error
        }
        
        self.current_evidence["database_queries"].append(query_evidence)
        logger.debug(f"Database query evidence added: {query}")
    
    def add_ui_screenshot(self, screenshot_path: str, description: str = "", page_url: str = ""):
        """
        Add UI screenshot evidence
        
        Args:
            screenshot_path: Path to screenshot file
            description: Description of the screenshot
            page_url: URL of the page captured
        """
        screenshot_info = {
            "timestamp": datetime.now().isoformat(),
            "path": screenshot_path,
            "description": description,
            "page_url": page_url
        }
        
        self.current_evidence["ui_screenshots"].append(screenshot_info)
        logger.debug(f"UI screenshot evidence added: {screenshot_path}")
    
    def save_evidence(self, output_format: str = "auto") -> Optional[Path]:
        """
        Save evidence to file
        
        Args:
            output_format: Output format - "auto" (detect from evidence type), "json", "docx", "pdf"
        
        Returns:
            Path to saved evidence file, or None if no evidence to save
        """
        if not self.current_evidence or not self.current_evidence.get("scenario"):
            return None
        
        # Create filename
        safe_feature = self._sanitize_filename(self.current_evidence["feature"])
        safe_scenario = self._sanitize_filename(self.current_evidence["scenario"])
        timestamp = self.current_evidence["timestamp"]
        
        # Create feature directory
        feature_dir = self.evidence_dir / safe_feature
        feature_dir.mkdir(parents=True, exist_ok=True)
        
        # Determine output format
        has_ui_screenshots = len(self.current_evidence.get("ui_screenshots", [])) > 0
        has_api_requests = len(self.current_evidence.get("api_requests", [])) > 0
        has_db_queries = len(self.current_evidence.get("database_queries", [])) > 0
        
        if output_format == "auto":
            # If UI screenshots exist, generate Word document
            if has_ui_screenshots:
                output_format = "docx"
            # Otherwise, use JSON for API/DB evidence
            elif has_api_requests or has_db_queries:
                output_format = "json"
            else:
                output_format = "json"
        
        # Save based on format
        if output_format == "docx" and has_ui_screenshots:
            return self._save_word_document(feature_dir, safe_feature, safe_scenario, timestamp)
        else:
            # Save as JSON
            filename = f"{safe_feature}_{safe_scenario}_{timestamp}.json"
            evidence_path = feature_dir / filename
            
            try:
                with open(evidence_path, 'w', encoding='utf-8') as f:
                    json.dump(self.current_evidence, f, indent=2, ensure_ascii=False)
                
                logger.info(f"Evidence saved to: {evidence_path}")
                return evidence_path
            except Exception as e:
                logger.error(f"Failed to save evidence: {str(e)}")
                return None
    
    def _save_word_document(self, feature_dir: Path, safe_feature: str, safe_scenario: str, timestamp: str) -> Optional[Path]:
        """
        Save UI evidence as Word document with screenshots
        
        Args:
            feature_dir: Feature directory
            safe_feature: Sanitized feature name
            safe_scenario: Sanitized scenario name
            timestamp: Timestamp string
            
        Returns:
            Path to saved Word document, or None if failed
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available. Cannot generate Word document.")
            return None
        
        try:
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'{self.current_evidence["feature"]} - {self.current_evidence["scenario"]}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_paragraph(f'Feature: {self.current_evidence["feature"]}')
            doc.add_paragraph(f'Scenario: {self.current_evidence["scenario"]}')
            doc.add_paragraph(f'Timestamp: {self.current_evidence["timestamp"]}')
            doc.add_paragraph('')  # Empty line
            
            # Add screenshots
            screenshots = self.current_evidence.get("ui_screenshots", [])
            if screenshots:
                doc.add_heading('Screenshots', level=1)
                
                for idx, screenshot_info in enumerate(screenshots, 1):
                    screenshot_path = Path(screenshot_info["path"])
                    
                    # Add screenshot section
                    doc.add_heading(f'Screenshot {idx}', level=2)
                    
                    # Add description if available
                    if screenshot_info.get("description"):
                        doc.add_paragraph(f'Description: {screenshot_info["description"]}')
                    
                    # Add page URL if available
                    if screenshot_info.get("page_url"):
                        doc.add_paragraph(f'Page URL: {screenshot_info["page_url"]}')
                    
                    # Add timestamp
                    if screenshot_info.get("timestamp"):
                        doc.add_paragraph(f'Time: {screenshot_info["timestamp"]}')
                    
                    # Add screenshot image
                    if screenshot_path.exists():
                        try:
                            paragraph = doc.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run()
                            run.add_picture(str(screenshot_path), width=Inches(6))
                            doc.add_paragraph('')  # Empty line after image
                        except Exception as e:
                            logger.warning(f"Failed to add screenshot {screenshot_path}: {str(e)}")
                            doc.add_paragraph(f'[Screenshot not available: {screenshot_path}]')
                    else:
                        doc.add_paragraph(f'[Screenshot file not found: {screenshot_path}]')
            
            # Add API requests if any
            api_requests = self.current_evidence.get("api_requests", [])
            if api_requests:
                doc.add_heading('API Requests', level=1)
                for idx, req in enumerate(api_requests, 1):
                    doc.add_heading(f'Request {idx}', level=2)
                    doc.add_paragraph(f'Method: {req["request"]["method"]}')
                    doc.add_paragraph(f'URL: {req["request"]["url"]}')
                    if req["request"].get("headers"):
                        doc.add_paragraph(f'Headers: {json.dumps(req["request"]["headers"], indent=2)}')
                    if req["request"].get("body"):
                        doc.add_paragraph(f'Body: {json.dumps(req["request"]["body"], indent=2)}')
                    doc.add_paragraph(f'Response Status: {req["response"]["status_code"]}')
                    if req["response"].get("body"):
                        doc.add_paragraph(f'Response: {json.dumps(req["response"]["body"], indent=2)}')
                    doc.add_paragraph('')
            
            # Add database queries if any
            db_queries = self.current_evidence.get("database_queries", [])
            if db_queries:
                doc.add_heading('Database Queries', level=1)
                for idx, query in enumerate(db_queries, 1):
                    doc.add_heading(f'Query {idx}', level=2)
                    doc.add_paragraph(f'SQL: {query["query"]}')
                    if query.get("error"):
                        doc.add_paragraph(f'Error: {query["error"]}')
                    else:
                        doc.add_paragraph(f'Rows: {query.get("row_count", 0)}')
                        if query.get("result"):
                            doc.add_paragraph(f'Result: {json.dumps(query["result"], indent=2)}')
                    doc.add_paragraph('')
            
            # Save document
            filename = f"{safe_feature}_{safe_scenario}_{timestamp}.docx"
            evidence_path = feature_dir / filename
            doc.save(str(evidence_path))
            
            logger.info(f"Evidence saved to Word document: {evidence_path}")
            
            # Clean up temporary screenshots directory
            try:
                if hasattr(self, 'screenshots_dir') and self.screenshots_dir.exists():
                    import shutil
                    shutil.rmtree(self.screenshots_dir)
            except Exception as e:
                logger.warning(f"Failed to clean up screenshots directory: {str(e)}")
            
            return evidence_path
            
        except Exception as e:
            logger.error(f"Failed to save Word document: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
    def _serialize_body(self, body: Any) -> Any:
        """
        Serialize request/response body for JSON storage
        
        Args:
            body: Body to serialize
            
        Returns:
            Serialized body
        """
        if body is None:
            return None
        
        # If already a dict or list, return as is
        if isinstance(body, (dict, list)):
            return body
        
        # If string, try to parse as JSON
        if isinstance(body, str):
            try:
                return json.loads(body)
            except (json.JSONDecodeError, ValueError):
                return body
        
        # For other types, convert to string
        return str(body)
    
    def _sanitize_filename(self, filename: str) -> str:
        """
        Sanitize filename for file system
        
        Args:
            filename: Original filename
            
        Returns:
            Sanitized filename
        """
        # Replace invalid characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        # Remove leading/trailing spaces and dots
        filename = filename.strip(' .')
        
        # Limit length
        if len(filename) > 100:
            filename = filename[:100]
        
        return filename

